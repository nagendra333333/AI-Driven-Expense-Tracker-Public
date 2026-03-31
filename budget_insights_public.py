"""
Budget Insights Generator
==========================
Reads  : Money_Manager_output/Money_Manager_*.tsv  (latest)
         Budget/Monthly_Budget_2026.xlsx
         OUTPUT_FOLDER/Full_statement_*.xlsx         (history for yearly cats)
Writes : Insights/Insights_<last_date>.docx

Yearly categories (prorated monthly in budget, accumulated across year):
  Cloth | Bike maintenance | Gas
  → Sum all Full_statement files YTD + current month
  → Compare against yearly budget (monthly × 12)
"""

import sys, json, re, os
from dotenv import load_dotenv
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from docx import Document
import pandas as pd

# ─── FOLDERS ─────────────────────────────────────────────────────────────────
SCRIPT_DIR       = Path(__file__).resolve().parent
MONEY_MGR_FOLDER = SCRIPT_DIR / "Money_Manager_output"
BUDGET_FOLDER    = SCRIPT_DIR / "Budget"
OUTPUT_FOLDER    = SCRIPT_DIR / "OUTPUT_FOLDER"
INSIGHTS_FOLDER  = SCRIPT_DIR / "Insights"

BUDGET_FILE      = BUDGET_FOLDER / "Monthly_Budget_2026.xlsx"

load_dotenv(SCRIPT_DIR / ".env")

# Categories that are budgeted monthly but tracked on a YEARLY basis
YEARLY_CATS = {"cloth", "bike maintenance", "gas"}

# ─── NVIDIA CLIENT ────────────────────────────────────────────────────────────
client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key=os.environ.get("NVIDIA_API_KEY", ""),  # Set in .env file
)

def ask_ai(prompt: str, max_tokens: int = 2048) -> str:
    try:
        completion = client.chat.completions.create(
            model="openai/gpt-oss-120b",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            top_p=1,
            max_tokens=max_tokens,
            stream=True,
        )
        raw = ""
        for chunk in completion:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta.content
            if delta:
                raw += delta
        return raw.strip()
    except Exception as e:
        print(f"  [WARN] AI API unavailable: {e}")
        return ""   # caller handles empty string gracefully

# ─── HELPERS ─────────────────────────────────────────────────────────────────
def find_latest_tsv() -> Path | None:
    files = sorted(MONEY_MGR_FOLDER.glob("Money_Manager_*.tsv"))
    return files[-1] if files else None

def load_tsv(path: Path) -> pd.DataFrame:
    df = pd.read_csv(str(path), sep="\t", dtype=str)
    df.columns = [c.strip() for c in df.columns]
    df["INR"] = pd.to_numeric(df["INR"], errors="coerce").fillna(0)
    return df

def load_budget() -> tuple[dict, dict, dict]:
    """Returns (expense_budget, investment_budget, income_budget) as dicts."""
    xf = pd.ExcelFile(str(BUDGET_FILE))

    # Expenses
    exp = pd.read_excel(str(BUDGET_FILE), sheet_name="Expenses", dtype=str)
    exp.columns = [c.strip() for c in exp.columns]
    exp_budget = {}
    for _, r in exp.iterrows():
        item = str(r.get("Item","")).strip()
        amt  = r.get("Amount","0")
        if item and item.lower() != "total" and item.lower() != "item":
            try:
                exp_budget[item] = float(str(amt).replace(",",""))
            except:
                pass

    # Investments
    inv = pd.read_excel(str(BUDGET_FILE), sheet_name="Investments", dtype=str)
    inv.columns = [c.strip() for c in inv.columns]
    inv_budget = {}
    for _, r in inv.iterrows():
        item = str(r.get("Item","")).strip()
        amt  = r.get("Amount","0")
        if item and item.lower() not in ("total","item"):
            try:
                inv_budget[item] = float(str(amt).replace(",",""))
            except:
                pass

    # Income
    inc = pd.read_excel(str(BUDGET_FILE), sheet_name="Income", dtype=str)
    inc.columns = [c.strip() for c in inc.columns]
    inc_budget = {}
    for _, r in inc.iterrows():
        item = str(r.get("Item","")).strip()
        amt  = r.get("Amount","0")
        if item and item.lower() not in ("total","item"):
            try:
                inc_budget[item] = float(str(amt).replace(",",""))
            except:
                pass

    return exp_budget, inv_budget, inc_budget

def normalize_cat(c: str) -> str:
    """Lowercase + strip for fuzzy budget matching."""
    return str(c).strip().lower()

def budget_key_for(cat: str, budget: dict) -> str | None:
    """Find the budget dict key that best matches a category name."""
    cl = normalize_cat(cat)
    # Exact (case-insensitive)
    for k in budget:
        if normalize_cat(k) == cl:
            return k
    # Partial
    for k in budget:
        kl = normalize_cat(k)
        if kl in cl or cl in kl:
            return k
    # Manual mappings
    ALIAS = {
        "electricity": "eb",
        "eb": "electricity",
        "dth + ott + net + mobile": "dth+ott+net+mobile",
        "dth+ott+net+mobile": "dth + ott + net + mobile",
        "home things": "misc",
        "miscellaneous": "misc",
        "mobile accessories": "misc",
        "medicine": "misc",
    }
    mapped = ALIAS.get(cl)
    if mapped:
        for k in budget:
            if normalize_cat(k) == mapped:
                return k
    return None

def load_ytd_yearly_cats(current_file: str) -> dict[str, float]:
    """
    Sum expenses for YEARLY_CATS across all Full_statement files EXCEPT current.
    Returns {cat_lower: total_spent_ytd}
    """
    totals: dict[str, float] = {c: 0.0 for c in YEARLY_CATS}
    for fp in sorted(OUTPUT_FOLDER.glob("Full_statement_*.xlsx")):
        if fp.name == current_file:
            continue
        try:
            df = pd.read_excel(str(fp), sheet_name="Transactions", dtype=str)
            df["Debit"] = pd.to_numeric(df["Debit"], errors="coerce").fillna(0)
            df["Expenses Category"] = df["Expenses Category"].fillna("")
            for cat in YEARLY_CATS:
                mask = df["Expenses Category"].str.strip().str.lower() == cat
                totals[cat] += df.loc[mask, "Debit"].sum()
        except Exception as e:
            print(f"  [history] Could not read {fp.name}: {e}")
    return totals

def latest_full_statement_name() -> str:
    files = sorted(OUTPUT_FOLDER.glob("Full_statement_*.xlsx"))
    return files[-1].name if files else ""

def fmt_inr(v: float) -> str:
    return f"₹{v:,.2f}"

def variance_symbol(actual: float, budget: float) -> str:
    if actual > budget * 1.1:
        return "🔴 OVER"
    elif actual < budget * 0.8:
        return "🟢 UNDER"
    else:
        return "🟡 ON TRACK"

# ─── MAIN ANALYSIS ────────────────────────────────────────────────────────────
def build_analysis(tsv_df: pd.DataFrame, exp_budget: dict,
                   inv_budget: dict, inc_budget: dict,
                   ytd_yearly: dict[str, float],
                   month_label: str) -> dict:
    """
    Returns a structured analysis dict used both for AI and for doc writing.
    """
    # ── Expense actuals from TSV ──────────────────────────────────────────────
    exp_rows = tsv_df[tsv_df["Income/Expense"] == "Expense"]
    expense_actual: dict[str, float] = {}
    for _, row in exp_rows.iterrows():
        cat = str(row.get("Category","")).strip()
        if cat:
            expense_actual[cat] = expense_actual.get(cat, 0) + float(row["INR"])

    # ── Income actuals ────────────────────────────────────────────────────────
    inc_rows = tsv_df[tsv_df["Income/Expense"] == "Income"]
    income_actual: dict[str, float] = {}
    for _, row in inc_rows.iterrows():
        cat = str(row.get("Category","")).strip()
        if cat:
            income_actual[cat] = income_actual.get(cat, 0) + float(row["INR"])

    # ── Budget comparisons ────────────────────────────────────────────────────
    budget_comparison = []
    for cat, actual in sorted(expense_actual.items(), key=lambda x: -x[1]):
        bkey = budget_key_for(cat, exp_budget)
        cat_l = normalize_cat(cat)

        if cat_l in YEARLY_CATS:
            # Yearly logic: compare YTD (prev months + this month) vs yearly budget
            ytd_prev     = ytd_yearly.get(cat_l, 0)
            ytd_total    = ytd_prev + actual
            monthly_bud  = exp_budget.get(bkey, 0) if bkey else 0
            yearly_bud   = monthly_bud * 12
            budget_comparison.append({
                "category"    : cat,
                "actual"      : actual,
                "budget"      : monthly_bud,
                "ytd_total"   : ytd_total,
                "yearly_budget": yearly_bud,
                "variance"    : actual - monthly_bud,
                "ytd_variance": ytd_total - yearly_bud,
                "status"      : variance_symbol(actual, monthly_bud),
                "is_yearly"   : True,
            })
        else:
            monthly_bud = exp_budget.get(bkey, 0) if bkey else 0
            budget_comparison.append({
                "category": cat,
                "actual"  : actual,
                "budget"  : monthly_bud,
                "variance": actual - monthly_bud,
                "status"  : variance_symbol(actual, monthly_bud) if monthly_bud else "⚪ UNBUDGETED",
                "is_yearly": False,
            })

    # ── Investment actuals vs budget ──────────────────────────────────────────
    inv_actual = expense_actual.get("Investment", 0)
    inv_total_budget = sum(inv_budget.values())
    inv_breakdown = []
    for k, v in inv_budget.items():
        inv_breakdown.append({"item": k, "budget": v})

    # ── Income summary ────────────────────────────────────────────────────────
    total_income_actual  = sum(income_actual.values())
    total_income_budget  = sum(inc_budget.values())
    total_expense_actual = sum(expense_actual.values())
    savings_actual       = total_income_actual - total_expense_actual

    return {
        "month_label"         : month_label,
        "budget_comparison"   : budget_comparison,
        "expense_actual"      : expense_actual,
        "income_actual"       : income_actual,
        "income_budget"       : inc_budget,
        "total_income_actual" : total_income_actual,
        "total_income_budget" : total_income_budget,
        "total_expense_actual": total_expense_actual,
        "inv_actual"          : inv_actual,
        "inv_total_budget"    : inv_total_budget,
        "inv_breakdown"       : inv_breakdown,
        "savings_actual"      : savings_actual,
        "exp_budget"          : exp_budget,
    }

# ─── AI INSIGHTS ──────────────────────────────────────────────────────────────
def generate_ai_insights(analysis: dict) -> dict:
    """Call AI to generate narrative insights. Returns dict of sections."""

    over_budget = [x for x in analysis["budget_comparison"]
                   if x["actual"] > x["budget"] and x["budget"] > 0]
    under_budget = [x for x in analysis["budget_comparison"]
                    if x["actual"] < x["budget"] * 0.8 and x["budget"] > 0]
    unbudgeted   = [x for x in analysis["budget_comparison"] if x["budget"] == 0]
    yearly_cats  = [x for x in analysis["budget_comparison"] if x.get("is_yearly")]

    prompt = f"""You are a personal finance advisor analyzing a monthly budget for an Indian family.

Month: {analysis['month_label']}

=== INCOME ===
Actual  : ₹{analysis['total_income_actual']:,.2f}
Budgeted: ₹{analysis['total_income_budget']:,.2f}
Income breakdown: {json.dumps({k: round(v,2) for k,v in analysis['income_actual'].items()})}

=== EXPENSE SUMMARY ===
Total actual   : ₹{analysis['total_expense_actual']:,.2f}
Total budgeted : ₹{sum(analysis['exp_budget'].values()):,.2f}

=== CATEGORIES OVER BUDGET ===
{json.dumps([{k: v for k,v in x.items() if k in ('category','actual','budget','variance')} for x in over_budget], indent=2)}

=== CATEGORIES UNDER BUDGET ===
{json.dumps([{k: v for k,v in x.items() if k in ('category','actual','budget','variance')} for x in under_budget], indent=2)}

=== UNBUDGETED EXPENSES ===
{json.dumps([{'category': x['category'], 'actual': x['actual']} for x in unbudgeted], indent=2)}

=== YEARLY TRACKED CATEGORIES (Cloth / Bike maintenance / Gas) ===
{json.dumps([{k: v for k,v in x.items() if k in ('category','actual','ytd_total','yearly_budget','ytd_variance')} for x in yearly_cats], indent=2)}

=== INVESTMENTS ===
Actual invested : ₹{analysis['inv_actual']:,.2f}
Total budgeted  : ₹{analysis['inv_total_budget']:,.2f}
Investment breakdown: {json.dumps({x['item']: x['budget'] for x in analysis['inv_breakdown']})}

=== NET SAVINGS ===
Savings this month: ₹{analysis['savings_actual']:,.2f}

Please generate a detailed financial insights report with these EXACT sections:
1. EXECUTIVE SUMMARY (2-3 sentences overview)
2. INCOME ANALYSIS (actual vs budget, any gaps)
3. EXPENSE OVERVIEW (total spend, key observations)
4. CATEGORY DEEP DIVE (for each over-budget category explain why it matters and impact)
5. INVESTMENT INSIGHTS (how investments compare to plan, which areas need attention)
6. YEARLY CATEGORY STATUS (cloth, bike maintenance, gas — yearly pacing analysis)
7. UNBUDGETED SPENDING (observations on unplanned expenses)
8. SAVINGS & NET POSITION (how savings compare to plan)
9. ACTIONABLE SUGGESTIONS (5-7 specific, practical suggestions for next month)

Write in a professional but friendly tone. Use Indian Rupees (₹). Be specific with numbers.
Format each section with its heading followed by content. No markdown, no bullet symbols in text.
Use numbered suggestions in section 9."""

    print("  Calling AI for insights…")
    raw = ask_ai(prompt, max_tokens=2048)

    # Parse into sections
    sections = {}
    current_key = None
    current_lines = []

    section_patterns = [
        "EXECUTIVE SUMMARY", "INCOME ANALYSIS", "EXPENSE OVERVIEW",
        "CATEGORY DEEP DIVE", "INVESTMENT INSIGHTS", "YEARLY CATEGORY STATUS",
        "UNBUDGETED SPENDING", "SAVINGS & NET POSITION", "ACTIONABLE SUGGESTIONS"
    ]

    for line in raw.split("\n"):
        matched = False
        for sp in section_patterns:
            if sp in line.upper():
                if current_key:
                    sections[current_key] = "\n".join(current_lines).strip()
                current_key   = sp
                current_lines = []
                matched       = True
                break
        if not matched and current_key:
            current_lines.append(line)

    if current_key:
        sections[current_key] = "\n".join(current_lines).strip()

    # Fallback: use full text if parsing failed
    if not sections:
        if raw:
            sections["FULL REPORT"] = raw
        else:
            sections = _rule_based_insights(analysis)

    return sections


def _rule_based_insights(analysis: dict) -> dict:
    """Generate basic insights without AI when API is unavailable."""
    over  = [x for x in analysis["budget_comparison"] if x["actual"] > x["budget"] > 0]
    under = [x for x in analysis["budget_comparison"] if x["actual"] < x["budget"]*0.8 and x["budget"] > 0]

    summary = (
        f"This report covers {analysis['month_label']}. "
        f"Total income was {fmt_inr(analysis['total_income_actual'])} against a budget of "
        f"{fmt_inr(analysis['total_income_budget'])}. "
        f"Total expenses were {fmt_inr(analysis['total_expense_actual'])}. "
        f"Net savings: {fmt_inr(analysis['savings_actual'])}."
    )

    cat_deep = ""
    for x in over:
        excess = x["actual"] - x["budget"]
        cat_deep += (f"{x['category']}: Spent {fmt_inr(x['actual'])} vs budget "
                     f"{fmt_inr(x['budget'])} — exceeded by {fmt_inr(excess)}. "
                     f"Status: {x['status']}\n")

    inv_text = (
        f"Total investment budget: {fmt_inr(analysis['inv_total_budget'])}. "
        f"Actual invested: {fmt_inr(analysis['inv_actual'])}. "
    )
    if analysis["inv_actual"] < analysis["inv_total_budget"] * 0.8:
        inv_text += "Investments are below plan — review SIP/equity contributions."
    elif analysis["inv_actual"] >= analysis["inv_total_budget"]:
        inv_text += "Investment targets met or exceeded — great discipline!"

    yearly_text = ""
    for x in analysis["budget_comparison"]:
        if x.get("is_yearly"):
            ytd   = x.get("ytd_total", x["actual"])
            yr_bud= x.get("yearly_budget", x["budget"]*12)
            yearly_text += (f"{x['category']}: YTD spend {fmt_inr(ytd)} vs yearly "
                            f"budget {fmt_inr(yr_bud)} "
                            f"({'over' if ytd > yr_bud else 'within'} limit).\n")

    suggestions = (
        "1. Review food spending — if over budget, plan weekly grocery lists.\n"
        "2. Track investments monthly against the plan sheet.\n"
        "3. For yearly categories (cloth, gas, bike), pace purchases evenly.\n"
        "4. Set aside unbudgeted expense categories into a miscellaneous fund.\n"
        "5. Review income sources — ensure salary credited on time.\n"
    )

    return {
        "EXECUTIVE SUMMARY"      : summary,
        "INCOME ANALYSIS"        : f"Income actual: {fmt_inr(analysis['total_income_actual'])}  |  Budget: {fmt_inr(analysis['total_income_budget'])}",
        "EXPENSE OVERVIEW"       : f"Total spent: {fmt_inr(analysis['total_expense_actual'])}  |  Budget: {fmt_inr(sum(analysis['exp_budget'].values()))}",
        "CATEGORY DEEP DIVE"     : cat_deep or "All categories within budget.",
        "INVESTMENT INSIGHTS"    : inv_text,
        "YEARLY CATEGORY STATUS" : yearly_text or "No yearly categories tracked this month.",
        "UNBUDGETED SPENDING"    : ", ".join(x["category"] for x in analysis["budget_comparison"] if x["budget"]==0) or "None",
        "SAVINGS & NET POSITION" : f"Net savings this month: {fmt_inr(analysis['savings_actual'])}",
        "ACTIONABLE SUGGESTIONS" : suggestions,
    }

# ─── WORD DOCUMENT ────────────────────────────────────────────────────────────
def write_docx(analysis: dict, ai_sections: dict, out_path: Path):
    """Write insights to a Word document using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Styles ────────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h1(text: str, color=(31,78,121)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.bold  = True
        run.font.size  = Pt(16)
        run.font.color.rgb = RGBColor(*color)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "%02X%02X%02X" % color)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def h2(text: str, color=(0, 112, 192)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.font.bold  = True
        run.font.size  = Pt(13)
        run.font.color.rgb = RGBColor(*color)
        return p

    def body(text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        return p

    def bullet(text: str, color=None):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_table_row(table, cols: list, bold=False, bg=None):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        row = table.add_row()
        for i, val in enumerate(cols):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.bold = bold
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bg:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd   = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), bg)
                tc_pr.append(shd)
        return row

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after  = Pt(6)
    tr = title.add_run(f"Monthly Budget Insights")
    tr.font.bold  = True
    tr.font.size  = Pt(22)
    tr.font.color.rgb = RGBColor(31,78,121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"{analysis['month_label']}  |  Generated {datetime.now().strftime('%d %b %Y')}")
    sr.font.size  = Pt(11)
    sr.font.color.rgb = RGBColor(128,128,128)
    sr.font.italic = True

    doc.add_paragraph()

    # ── Summary Scorecard ─────────────────────────────────────────────────────
    h1("FINANCIAL SCORECARD")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, label in enumerate(["Metric", "Budget", "Actual"]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)

    scorecard_rows = [
        ("Total Income",     fmt_inr(analysis['total_income_budget']),  fmt_inr(analysis['total_income_actual'])),
        ("Total Expenses",   fmt_inr(sum(analysis['exp_budget'].values())), fmt_inr(analysis['total_expense_actual'])),
        ("Investments",      fmt_inr(analysis['inv_total_budget']),       fmt_inr(analysis['inv_actual'])),
        ("Net Savings",      "₹0",                                        fmt_inr(analysis['savings_actual'])),
    ]

    row_colors = ["DDEEFF", "DDEEFF", "DDEEFF", "D5E8D4"]
    for (label, bud, act), bg in zip(scorecard_rows, row_colors):
        add_table_row(tbl, [label, bud, act], bg=bg)

    doc.add_paragraph()

    # ── AI Sections ───────────────────────────────────────────────────────────
    section_order = [
        ("EXECUTIVE SUMMARY",      (31,78,121)),
        ("INCOME ANALYSIS",        (0,112,192)),
        ("EXPENSE OVERVIEW",       (192,80,77)),
        ("CATEGORY DEEP DIVE",     (192,80,77)),
        ("INVESTMENT INSIGHTS",    (68,114,196)),
        ("YEARLY CATEGORY STATUS", (84,130,53)),
        ("UNBUDGETED SPENDING",    (204,102,0)),
        ("SAVINGS & NET POSITION", (0,128,0)),
        ("ACTIONABLE SUGGESTIONS", (31,78,121)),
        ("FULL REPORT",            (31,78,121)),
    ]

    for section_key, color in section_order:
        text = ai_sections.get(section_key, "").strip()
        if not text:
            continue
        h1(section_key.title(), color=color)
        for para in text.split("\n"):
            para = para.strip()
            if not para:
                continue
            # Numbered suggestions
            if re.match(r"^\d+[\.\)]\s", para):
                bullet(para)
            else:
                body(para)

    # ── Budget Comparison Table ────────────────────────────────────────────────
    h1("DETAILED BUDGET vs ACTUAL", color=(31,78,121))

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Category", "Budget", "Actual", "Status"]):
        tbl2.rows[0].cells[i].text = h
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for item in sorted(analysis["budget_comparison"], key=lambda x: -x["actual"]):
        status = item.get("status","")
        bg = "FFE6E6" if "OVER" in status else ("E6FFE6" if "UNDER" in status else "FFFDE6")
        note = ""
        if item.get("is_yearly"):
            note = f" [YTD: {fmt_inr(item.get('ytd_total',0))} / Yearly: {fmt_inr(item.get('yearly_budget',0))}]"
        add_table_row(tbl2, [
            item["category"] + note,
            fmt_inr(item["budget"]),
            fmt_inr(item["actual"]),
            status,
        ], bg=bg)

    doc.add_paragraph()

    # ── Investment Breakdown ──────────────────────────────────────────────────
    h1("INVESTMENT PLAN vs ACTUAL", color=(68,114,196))
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.style = "Table Grid"
    tbl3.rows[0].cells[0].text = "Investment Type"
    tbl3.rows[0].cells[1].text = "Monthly Budget"
    for cell in tbl3.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    for item in analysis["inv_breakdown"]:
        add_table_row(tbl3, [item["item"], fmt_inr(item["budget"])], bg="EEF3FB")

    add_table_row(tbl3, ["TOTAL BUDGETED", fmt_inr(analysis["inv_total_budget"])],
                  bold=True, bg="DDEEFF")
    add_table_row(tbl3, ["ACTUAL INVESTED", fmt_inr(analysis["inv_actual"])],
                  bold=True, bg="D5E8D4")

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Generated by AI Budget Insights  |  Powered by NVIDIA GPT-4o 120B")
    fr.font.size   = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(150,150,150)

    doc.save(str(out_path))
    print(f"  Document saved → {out_path.name}")

# ─── ENTRY POINT ──────────────────────────────────────────────────────────────
def main():
    print("=" * 62)
    print("  Budget Insights Generator")
    print("=" * 62)

    # Ensure folders
    INSIGHTS_FOLDER.mkdir(exist_ok=True)

    # ── Find latest Money Manager TSV ─────────────────────────────────────────
    tsv_path = find_latest_tsv()
    if tsv_path is None:
        print(f"[ERROR] No Money_Manager_*.tsv in {MONEY_MGR_FOLDER}")
        sys.exit(1)
    print(f"  TSV    : {tsv_path.name}")

    # Extract date from filename for output
    stmt_date = tsv_path.stem.replace("Money_Manager_", "")   # e.g. 2026-03-23
    month_label = datetime.strptime(stmt_date, "%Y-%m-%d").strftime("%B %Y")
    print(f"  Month  : {month_label}")

    # ── Load data ─────────────────────────────────────────────────────────────
    tsv_df = load_tsv(tsv_path)
    print(f"  TSV rows: {len(tsv_df)}")

    exp_budget, inv_budget, inc_budget = load_budget()
    print(f"  Budget loaded: {len(exp_budget)} expense categories, "
          f"{len(inv_budget)} investment items")

    # ── Yearly category history ───────────────────────────────────────────────
    latest_fs = latest_full_statement_name()
    ytd_yearly = load_ytd_yearly_cats(latest_fs)
    if any(v > 0 for v in ytd_yearly.values()):
        print(f"  YTD history: {', '.join(f'{k}=₹{v:,.0f}' for k,v in ytd_yearly.items() if v)}")
    else:
        print("  YTD history: no prior months found (first month of year)")

    # ── Build analysis ────────────────────────────────────────────────────────
    analysis = build_analysis(tsv_df, exp_budget, inv_budget, inc_budget,
                              ytd_yearly, month_label)

    over  = [x for x in analysis["budget_comparison"] if x["actual"] > x["budget"] > 0]
    under = [x for x in analysis["budget_comparison"] if x["actual"] < x["budget"] * 0.8 and x["budget"] > 0]
    print(f"\n  Over budget  : {len(over)} categories")
    print(f"  Under budget : {len(under)} categories")
    print(f"  Net savings  : {fmt_inr(analysis['savings_actual'])}")

    # ── AI Insights ───────────────────────────────────────────────────────────
    print(f"\n{'─'*62}")
    print("  Generating AI insights…")
    print(f"{'─'*62}")
    ai_sections = generate_ai_insights(analysis)
    print(f"  AI sections generated: {list(ai_sections.keys())}")

    # ── Write Word document ───────────────────────────────────────────────────
    out_path = INSIGHTS_FOLDER / f"Insights_{stmt_date}.docx"
    print(f"\n  Writing Word document…")
    write_docx(analysis, ai_sections, out_path)

    print(f"\n{'═'*62}")
    print(f"  ✅  Saved → {out_path}")
    print(f"{'═'*62}")

if __name__ == "__main__":
    main()
